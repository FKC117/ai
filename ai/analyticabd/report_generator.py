from django.shortcuts import render, redirect
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from .models import UserDataset, DatasetVariable, AnalysisSession, AnalysisInteraction, UserPreference, AnalysisHistory, UserWarningPreference, ReportDocument, ReportSection, ChatMessage
import pandas as pd
import numpy as np
import json
import os
from datetime import datetime
from io import BytesIO
import re
from bs4 import BeautifulSoup, NavigableString
import requests
from urllib.parse import urljoin
from django.conf import settings

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def is_table_line(line):
    """Check if a line is part of a markdown table"""
    line_stripped = line.strip()
    # Check for pipe-separated tables
    if line_stripped.startswith('|') and line_stripped.endswith('|'):
        return True
    # Check for separator rows (dashes)
    if line_stripped.replace('-', '').replace(':', '').replace('|', '').strip() == '' and '|' in line_stripped:
        return True
    # Check for tab-separated tables (common in AI responses)
    if '\t' in line_stripped and len(line_stripped.split('\t')) > 1:
        return True
    return False


def render_table_from_metadata(doc, table_data):
    """Render a table from metadata with proper styling"""
    try:
        if table_data.get('type') == 'markdown':
            # Handle markdown table
            markdown_text = table_data.get('markdown', '')
            table_lines = markdown_text.split('\n')
            parsed_rows = []
            
            for line in table_lines:
                line_stripped = line.strip()
                if line_stripped.startswith('|') and line_stripped.endswith('|'):
                    # Remove leading/trailing | and split by |
                    cells = [cell.strip() for cell in line.strip('|').split('|')]
                    parsed_rows.append(cells)
                elif '\t' in line_stripped:
                    # Handle tab-separated tables
                    cells = [cell.strip() for cell in line_stripped.split('\t')]
                    parsed_rows.append(cells)
            
            if parsed_rows and len(parsed_rows) > 1:  # At least header + data
                # Skip separator row (dashes)
                data_rows = [row for row in parsed_rows if not all(cell.replace('-', '').replace(':', '').strip() == '' for cell in row)]
                
                if data_rows:
                    doc_table = doc.add_table(rows=len(data_rows), cols=len(data_rows[0]))
                    doc_table.style = 'Table Grid'
                    
                    for r_idx, row_cells in enumerate(data_rows):
                        for c_idx, cell_text in enumerate(row_cells):
                            if c_idx < len(doc_table.rows[r_idx].cells):
                                doc_cell = doc_table.rows[r_idx].cells[c_idx]
                                doc_cell.text = cell_text
                                
                                # Style the cell
                                for paragraph in doc_cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Inter'
                                        run.font.size = Pt(10)
                                        if r_idx == 0:  # Header row
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                            doc_cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                                            doc_cell._tc.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), '1a365d')  # Primary color
                                        else:
                                            run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
                                            if r_idx % 2 == 1:  # Alternating rows
                                                doc_cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                                                doc_cell._tc.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), 'f7fafc')  # --color-background
                    
                    # Auto-fit column widths
                    for column in doc_table.columns:
                        column.width = Inches(1.5)
                    
                    doc.add_paragraph('')  # Add space after table
                    print(f"✅ Added markdown table from metadata successfully")
                    return True
                else:
                    print(f"⚠️ No valid data rows found in markdown table")
            else:
                print(f"⚠️ Invalid markdown table format")
        
        elif table_data.get('type') == 'html':
            # Handle HTML table with enhanced styling
            soup = BeautifulSoup(table_data['html'], 'html.parser')
            table = soup.find('table')
            
            if table:
                # Create a new table in the document
                rows = table.find_all('tr')
                if rows:
                    # Determine max columns across all rows
                    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                    doc_table = doc.add_table(rows=len(rows), cols=max_cols)
                    doc_table.style = 'Table Grid'
                    
                    for r_idx, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for c_idx, cell in enumerate(cells):
                            if c_idx < len(doc_table.rows[r_idx].cells):
                                doc_cell = doc_table.rows[r_idx].cells[c_idx]
                                doc_cell.text = cell.get_text(strip=True)
                                
                                # Enhanced styling
                                for paragraph in doc_cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Inter'
                                        run.font.size = Pt(9)  # Slightly smaller for better fit
                                        
                                        # Check if this is a header cell (th tag or first row)
                                        is_header = cell.name == 'th' or r_idx == 0
                                        
                                        if is_header:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                            doc_cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                                            doc_cell._tc.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), '1a365d')  # Primary color
                                        else:
                                            run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
                                            # Alternating row colors
                                            if r_idx % 2 == 1:
                                                doc_cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                                                doc_cell._tc.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), 'f7fafc')  # Light background
                    
                    # Auto-fit column widths based on content
                    for column in doc_table.columns:
                        column.width = Inches(1.2)  # Slightly narrower for better fit
                    
                    doc.add_paragraph('')  # Add space after table
                    print(f"✅ Added enhanced HTML table from metadata successfully")
                    return True
                else:
                    print(f"⚠️ No rows found in table HTML")
            else:
                print(f"⚠️ No table tag found in HTML")
        else:
            print(f"⚠️ Unknown table type: {table_data.get('type')}")
    except Exception as e:
        print(f"❌ Error rendering table from metadata: {e}")
    
    return False


def render_images_from_metadata(doc, images, base_url=None, cookies=None):
    """Render images from metadata (supports type=url) with optional captions."""
    try:
        if not images:
            return 0
        added = 0
        for img in images:
            src = img.get('src')
            if not src:
                continue
            try:
                # Normalize to absolute URL if a base_url is given and src is relative
                src_url = src
                if base_url and isinstance(src, str) and (src.startswith('/') or not src.lower().startswith('http')):
                    # Fix common typos in path keys
                    fixed = src.replace('analyysis', 'analysis')
                    if not fixed.startswith('/'):
                        fixed = '/' + fixed
                    src_url = urljoin(base_url, fixed)

                resp = requests.get(src_url, timeout=10, cookies=cookies or {})
                if resp.status_code == 200:
                    try:
                        pic_stream = BytesIO(resp.content)
                        doc.add_picture(pic_stream, width=Inches(6.0))
                        caption = img.get('title') or 'Figure'
                        cap = doc.add_paragraph(caption)
                        if cap.runs:
                            cap.runs[0].font.name = 'Inter'
                            cap.runs[0].font.size = Pt(9)
                        added += 1
                        doc.add_paragraph('')
                    except Exception as pic_e:
                        print(f"⚠️ Unable to insert image from {src_url}: {pic_e}")
                else:
                    print(f"⚠️ Image fetch failed: {src_url} status {resp.status_code}")
            except Exception as e:
                print(f"⚠️ Image fetch exception for {src_url if 'src_url' in locals() else src}: {e}")
        return added
    except Exception as e:
        print(f"❌ Error rendering images: {e}")
        return 0


def process_paragraph(doc, paragraph_text):
    """Process accumulated paragraph text with proper styling"""
    if paragraph_text:
        if paragraph_text.startswith('##'):
            subheading = doc.add_heading(paragraph_text.replace('##', '').strip(), level=3)
            for run in subheading.runs:
                run.font.name = 'Inter'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(56, 178, 172)  # --color-secondary
        elif paragraph_text.startswith('- **'):
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            run = p.add_run(paragraph_text.replace('- **', '').replace('**', ''))
            run.font.name = 'Inter'
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 54, 93)  # --color-primary
        elif paragraph_text.startswith('- '):
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            run = p.add_run(paragraph_text.replace('- ', ''))
            run.font.name = 'Inter'
            run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
        elif paragraph_text.strip():
            p = doc.add_paragraph(paragraph_text)
            for run in p.runs:
                run.font.name = 'Inter'
                run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary


def extract_html_tables_from_content(content):
    """Extract HTML tables from AI response content and convert them to Word format"""
    tables = []
    
    try:
        # Use BeautifulSoup to parse HTML content
        soup = BeautifulSoup(content, 'html.parser')
        
        # Find all table elements
        html_tables = soup.find_all('table')
        
        for i, table in enumerate(html_tables):
            # Extract table data
            rows = table.find_all('tr')
            if rows:
                # Determine max columns across all rows
                max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                
                # Create table data structure
                table_data = {
                    'type': 'html',
                    'html': str(table),
                    'title': f'Table {i+1}',
                    'rows': len(rows),
                    'columns': max_cols
                }
                tables.append(table_data)
                print(f"📊 Extracted HTML table {i+1} with {len(rows)} rows and {max_cols} columns")
        
        return tables
        
    except Exception as e:
        print(f"❌ Error extracting HTML tables: {e}")
        return []


def create_document_header(doc, dataset):
    """Create the document header with DataFlow Analytics branding"""
    from docx.shared import RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn

    # Title with DataFlow Analytics branding
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("DataFlow Analytics Report")
    title_run.font.name = 'Inter'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 54, 93)  # --color-primary
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle_paragraph = doc.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run("Advanced Data Analysis & Insights")
    subtitle_run.font.name = 'Inter'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(113, 128, 150)  # --color-text-secondary
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a decorative line
    doc.add_paragraph('')

    # Dataset Information Card
    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dataset name with primary color
    dataset_name_run = info_paragraph.add_run(f"Dataset: {dataset.name}")
    dataset_name_run.font.name = 'Inter'
    dataset_name_run.font.size = Pt(16)
    dataset_name_run.font.bold = True
    dataset_name_run.font.color.rgb = RGBColor(26, 54, 93)  # --color-primary
    
    info_paragraph.add_run("\n")
    
    # Dataset stats with secondary color
    stats_run = info_paragraph.add_run(f"Rows: {dataset.rows:,} | Columns: {dataset.columns} | Size: {dataset.file_size or 'N/A'}")
    stats_run.font.name = 'Inter'
    stats_run.font.size = Pt(12)
    stats_run.font.color.rgb = RGBColor(56, 178, 172)  # --color-secondary
    
    info_paragraph.add_run("\n")
    
    # Upload and generation dates
    date_run = info_paragraph.add_run(f"Uploaded: {dataset.uploaded_at.strftime('%B %d, %Y at %I:%M %p')}")
    date_run.font.name = 'Inter'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(113, 128, 150)  # --color-text-secondary
    
    date_run = info_paragraph.add_run(f" | Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_run.font.name = 'Inter'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(113, 128, 150)  # --color-text-secondary

    doc.add_paragraph('')


def create_document_footer(doc):
    """Create the document footer with DataFlow branding"""
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Footer with DataFlow branding
    footer_paragraph = doc.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run = footer_paragraph.add_run("Generated by DataFlow Analytics")
    footer_run.font.name = 'Inter'
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(113, 128, 150)  # --color-text-secondary
    
    footer_paragraph.add_run(" | ")
    
    footer_run = footer_paragraph.add_run("Advanced AI-Powered Data Analysis")
    footer_run.font.name = 'Inter'
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(56, 178, 172)  # --color-secondary


def process_ai_response_section(doc, content, tables_data):
    """Process AI response section with exact order preservation"""
    try:
        import re
        
        # Parse content with BeautifulSoup to preserve exact order
        soup = BeautifulSoup(content, 'html.parser')
        table_index = 0  # Track which table from metadata we're on
        
        # Process each top-level element in order
        for element in soup.contents:
            if element.name == 'table':
                # This is an HTML table - render it directly
                table_data = {
                    'type': 'html',
                    'html': str(element),
                    'title': 'Table'
                }
                render_table_from_metadata(doc, table_data)
            elif isinstance(element, NavigableString):
                # This is plain text - process it line by line
                text_content = str(element)
                lines = text_content.splitlines()
                current_paragraph = []
                
                for line in lines:
                    line_stripped = line.strip()
                    
                    # Check if this line is part of a markdown table
                    if is_table_line(line):
                        # Process any accumulated text first
                        if current_paragraph:
                            paragraph_text = '\n'.join(current_paragraph).strip()
                            process_paragraph(doc, paragraph_text)
                            current_paragraph = []
                        
                        # Start collecting table lines
                        table_lines = [line]
                        in_table = True
                        
                        # Continue reading lines until we're out of the table
                        while in_table and lines:
                            next_line = lines[0] if lines else ""
                            if is_table_line(next_line):
                                table_lines.append(next_line)
                                lines.pop(0)  # Remove the line we just processed
                            else:
                                in_table = False
                        
                        # Try to render the collected table
                        if table_index < len(tables_data):
                            table_rendered = render_table_from_metadata(doc, tables_data[table_index])
                            if table_rendered:
                                table_index += 1
                        else:
                            # Fallback: render as plain text if no metadata table available
                            for table_line in table_lines:
                                p = doc.add_paragraph(table_line)
                                for run in p.runs:
                                    run.font.name = 'Inter'
                                    run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
                    else:
                        # If this is a blank line, process accumulated paragraph
                        if not line_stripped and current_paragraph:
                            paragraph_text = '\n'.join(current_paragraph).strip()
                            process_paragraph(doc, paragraph_text)
                            current_paragraph = []
                        else:
                            # Add line to current paragraph
                            current_paragraph.append(line)
                
                # Process any remaining paragraph
                if current_paragraph:
                    paragraph_text = '\n'.join(current_paragraph).strip()
                    process_paragraph(doc, paragraph_text)
            else:
                # Handle other HTML elements (if any) as text
                text_content = element.get_text()
                if text_content.strip():
                    p = doc.add_paragraph(text_content)
                    for run in p.runs:
                        run.font.name = 'Inter'
                        run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
    
    except Exception as e:
        print(f"⚠️ Error processing AI response content: {e}")
        # Fallback to simple text processing
        lines = content.splitlines()
        for line in lines:
            if line.strip():
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = 'Inter'
                    run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary


def process_non_ai_section(doc, content):
    """Process non-AI response sections"""
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn

    # Filter out table lines from regular content
    non_table_lines = []
    filtered_table_rows = []  # Define the missing variable
    for line in content.splitlines():
        if is_table_line(line):
            filtered_table_rows.append(line)  # Collect table lines
        elif line.strip():
            non_table_lines.append(line)
    
    if filtered_table_rows:
        parsed_rows = []
        for row in filtered_table_rows:
            cells = [c.strip() for c in row.strip('|').split('|')]
            parsed_rows.append(cells)
        if parsed_rows:
            table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
            table.style = 'Table Grid'
            
            for r_idx, row_cells in enumerate(parsed_rows):
                for c_idx, cell_text in enumerate(row_cells):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = str(cell_text)
                    
                    if r_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Inter'
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                run.font.size = Pt(11)
                        cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                        cell._tc.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), '1a365d')  # Primary color
                    else:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Inter'
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
                        
                        if r_idx % 2 == 1:
                            cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
                            cell._tc.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), 'f7fafc')  # --color-background
            
            for column in table.columns:
                column.width = Inches(1.5)
        
        if non_table_lines:
            for line in non_table_lines:
                line = line.strip()
                if line.startswith('##'):
                    subheading = doc.add_heading(line.replace('##', '').strip(), level=3)
                    for run in subheading.runs:
                        run.font.name = 'Inter'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(56, 178, 172)  # --color-secondary
                elif line.startswith('- **'):
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    run = p.add_run(line.replace('- **', '').replace('**', ''))
                    run.font.name = 'Inter'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)  # --color-primary
                elif line.startswith('- '):
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    run = p.add_run(line.replace('- ', ''))
                    run.font.name = 'Inter'
                    run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
                elif line.strip():
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = 'Inter'
                        run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
    else:
        # Process non-table content only
        paragraphs = '\n'.join(non_table_lines).split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if para.startswith('##'):
                subheading = doc.add_heading(para.replace('##', '').strip(), level=3)
                for run in subheading.runs:
                    run.font.name = 'Inter'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(56, 178, 172)  # --color-secondary
            elif para.startswith('- **'):
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                run = p.add_run(para.replace('- **', '').replace('**', ''))
                run.font.name = 'Inter'
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 54, 93)  # --color-primary
            elif para.startswith('- '):
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                run = p.add_run(para.replace('- ', ''))
                run.font.name = 'Inter'
                run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary
            elif para.strip():
                p = doc.add_paragraph(para)
                for run in p.runs:
                    run.font.name = 'Inter'
                    run.font.color.rgb = RGBColor(45, 55, 72)  # --color-text-primary

    doc.add_paragraph('')


def generate_report_document(request, dataset_id, session_id):
    """Generate the complete report document"""
    try:
        if not DOCX_AVAILABLE:
            return JsonResponse({'error': 'python-docx is not installed on the server'}, status=500)

        dataset = UserDataset.objects.get(id=dataset_id, user=request.user)
        # Resolve session: accept numeric ID or fallback to active session
        try:
            session = AnalysisSession.objects.get(id=int(session_id), user=request.user, dataset=dataset)
        except Exception:
            session = AnalysisSession.objects.filter(user=request.user, dataset=dataset, is_active=True).order_by('-updated_at').first()
            if not session:
                return JsonResponse({'error': 'No active session found for this dataset'}, status=404)
        document_obj = ReportDocument.objects.filter(user=request.user, dataset=dataset, session=session).first()

        if not document_obj:
            return JsonResponse({'error': 'No report content to download yet'}, status=404)

        # Build DOCX
        doc = Document()

        # Create document header
        create_document_header(doc, dataset)

        # Sections - smart ordering by type first, then by order/created_at
        type_priority = {
            'summary_statistics': 10,
            'distributions': 20,
            'correlation': 30,
            'outliers': 40,
            'data_quality': 50,
            'visualization': 60,
            'ai_response': 90,
        }

        sections = list(document_obj.sections.all())
        sections.sort(key=lambda s: (type_priority.get(s.section_type, 80), s.order, s.created_at))

        # Insert a top-level EDA heading once if any EDA sections exist
        any_eda = any((s.metadata or {}).get('eda_subsection') for s in sections)
        eda_heading_added = False

        for section in sections:
            # Section header with DataFlow styling
            # If this is an EDA subsection and we haven't added the EDA main heading yet
            is_eda_sub = (section.metadata or {}).get('eda_subsection') is not None
            if is_eda_sub and not eda_heading_added:
                main = doc.add_heading('Exploratory Data Analysis (EDA)', level=1)
                for run in main.runs:
                    run.font.name = 'Inter'
                    run.font.size = Pt(22)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
                eda_heading_added = True

            heading_level = 3 if is_eda_sub else 2
            heading = doc.add_heading(section.title, level=heading_level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Style the heading
            for run in heading.runs:
                run.font.name = 'Inter'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 54, 93)  # --color-primary

            content = section.content or ''
            
            # Check if section has report data (tables/images from AI responses)
            report_data = section.metadata.get('report_data', {}) if section.metadata else {}
            tables_data = report_data.get('tables', [])
            images_data = report_data.get('images', [])
            
            print(f"📄 Processing section '{section.title}' with {len(tables_data)} tables and {len(images_data)} images")
            
            # For AI responses, preserve the exact order of content (text, tables, text)
            if section.section_type == 'ai_response':
                process_ai_response_section(doc, content, tables_data)
            else:
                # For non-AI responses, render images first (charts), then text, then tables
                if images_data:
                    base_url = request.build_absolute_uri('/')
                    # Pass session cookie for authenticated endpoints
                    cookies = {}
                    session_cookie_name = getattr(settings, 'SESSION_COOKIE_NAME', 'sessionid')
                    if request.COOKIES.get(session_cookie_name):
                        cookies[session_cookie_name] = request.COOKIES.get(session_cookie_name)
                    _img_count = render_images_from_metadata(doc, images_data, base_url=base_url, cookies=cookies)
                    print(f"🖼️ Rendered {_img_count} images for section '{section.title}'")
                process_non_ai_section(doc, content)
                if tables_data:
                    for tbl in tables_data:
                        render_table_from_metadata(doc, tbl)

        # Create document footer
        create_document_footer(doc)

        # Stream the document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"DataFlow_Report_{dataset.id}_{session.id}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    except UserDataset.DoesNotExist:
        return JsonResponse({'error': 'Dataset not found'}, status=404)
    except AnalysisSession.DoesNotExist:
        return JsonResponse({'error': 'Session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': f'Error generating report: {str(e)}'}, status=500)
